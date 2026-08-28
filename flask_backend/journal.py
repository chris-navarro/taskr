from calendar import monthrange
from datetime import datetime, timedelta
from io import StringIO
import csv

from flask import Blueprint, Response, render_template, request
from flask_login import current_user, login_required

from .models.task import Task
from .models.task_update import TaskUpdate

journal = Blueprint("journal", __name__, url_prefix="/journal")


def period_bounds(period, reference=None):
    reference = reference or datetime.utcnow()
    if period == "monthly":
        start = reference.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        end = start.replace(day=monthrange(start.year, start.month)[1], hour=23, minute=59, second=59)
    elif period == "mid-year":
        start_month = 1 if reference.month <= 6 else 7
        start = reference.replace(month=start_month, day=1, hour=0, minute=0, second=0, microsecond=0)
        end_month = start_month + 5
        end = start.replace(month=end_month, day=monthrange(start.year, end_month)[1], hour=23, minute=59, second=59)
    elif period == "year-end":
        start = reference.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
        end = reference.replace(month=12, day=31, hour=23, minute=59, second=59, microsecond=0)
    else:
        start = (reference - timedelta(days=reference.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
        end = start + timedelta(days=6, hours=23, minutes=59, seconds=59)
    return start, end


def build_report(period, custom_start=None, custom_end=None, category=None, status=None, project=None):
    start, end = (custom_start, custom_end) if custom_start and custom_end else period_bounds(period)
    tasks = Task.all(owner_id=current_user.id)
    entries = []
    updates = []
    for task in tasks:
        if category and task.get("category") != category:
            continue
        if project and task.get("project_id") != project:
            continue
        task_updates = TaskUpdate.between(task["_id"], start, end)
        if status:
            task_updates = [update for update in task_updates if (update.get("status") or "Updated") == status]
        if task_updates:
            entries.append({"task": task, "updates": task_updates})
            updates.extend(task_updates)

    status_counts = {}
    category_counts = {}
    daily_hours = {}
    planned_hours = sum(float(entry["task"].get("estimated_hours", 0) or 0) for entry in entries)
    for update in updates:
        status = update.get("status") or "Updated"
        status_counts[status] = status_counts.get(status, 0) + 1
        day = update.get("worked_at")
        if day:
            day_label = day.strftime("%b %d")
            daily_hours[day_label] = daily_hours.get(day_label, 0) + float(update.get("hours_worked", 0) or 0)
    for entry in entries:
        category = entry["task"].get("category") or "Uncategorised"
        category_counts[category] = category_counts.get(category, 0) + len(entry["updates"])

    duration = end - start
    previous_start = start - duration - timedelta(seconds=1)
    previous_end = start - timedelta(seconds=1)
    previous_updates = []
    for task in tasks:
        if ((category and task.get("category") != category) or
                (project and task.get("project_id") != project)):
            continue
        task_previous_updates = TaskUpdate.between(task["_id"], previous_start, previous_end)
        if status:
            task_previous_updates = [update for update in task_previous_updates if (update.get("status") or "Updated") == status]
        previous_updates.extend(task_previous_updates)
    previous_hours = sum(float(update.get("hours_worked", 0) or 0) for update in previous_updates)
    actual_hours = sum(float(update.get("hours_worked", 0) or 0) for update in updates)
    completed_updates = sum((update.get("status") or "").lower() == "completed" for update in updates)
    scorecards = {
        "completion_rate": round(completed_updates / len(updates) * 100) if updates else 0,
        "accomplishment_rate": round(sum(bool(update.get("accomplishment")) for update in updates) / len(updates) * 100) if updates else 0,
        "planned_utilization": round(actual_hours / planned_hours * 100) if planned_hours else 0,
        "active_days": len(daily_hours),
    }

    return {
        "period": period,
        "period_label": f"{start.strftime('%b %d, %Y')} - {end.strftime('%b %d, %Y')}",
        "start": start,
        "end": end,
        "entries": entries,
        "stats": {
            "updates": len(updates),
            "hours": sum(float(update.get("hours_worked", 0) or 0) for update in updates),
            "accomplishments": sum(bool(update.get("accomplishment")) for update in updates),
            "completed": completed_updates,
            "planned_hours": planned_hours,
            "actual_hours": actual_hours,
            "tasks": len(entries),
        },
        "status_counts": sorted(status_counts.items()),
        "category_counts": sorted(category_counts.items()),
        "daily_hours": list(daily_hours.items()),
        "filters": {"category": category or "", "status": status or "", "project": project or ""},
        "comparison": {"updates": len(updates) - len(previous_updates), "hours": actual_hours - previous_hours},
        "scorecards": scorecards,
    }


def available_filters():
    tasks = Task.all(owner_id=current_user.id)
    return {
        "categories": sorted({task.get("category") for task in tasks if task.get("category")}),
        "projects": sorted({task.get("project_id") for task in tasks if task.get("project_id")}),
        "statuses": sorted({task.get("status") for task in tasks if task.get("status")}),
    }


@journal.route("/")
@login_required
def report():
    period = request.args.get("period", "weekly")
    custom_start = custom_end = None
    if period == "custom":
        try:
            custom_start = datetime.strptime(request.args["start"], "%Y-%m-%d")
            custom_end = datetime.strptime(request.args["end"], "%Y-%m-%d").replace(hour=23, minute=59, second=59)
            if custom_end < custom_start:
                raise ValueError
        except (KeyError, ValueError):
            period = "weekly"
    if period not in {"weekly", "monthly", "mid-year", "year-end", "custom"}:
        period = "weekly"
    filters = available_filters()
    selected = {key: request.args.get(key, "") for key in ("category", "status", "project")}
    return render_template("journal/report.html", report=build_report(period, custom_start, custom_end, **selected), filters=filters)


@journal.route("/export.csv")
@login_required
def export_csv():
    period = request.args.get("period", "weekly")
    custom_start = custom_end = None
    if period == "custom":
        try:
            custom_start = datetime.strptime(request.args["start"], "%Y-%m-%d")
            custom_end = datetime.strptime(request.args["end"], "%Y-%m-%d").replace(hour=23, minute=59, second=59)
        except (KeyError, ValueError):
            period = "weekly"
    report_data = build_report(period, custom_start, custom_end, request.args.get("category"), request.args.get("status"), request.args.get("project"))
    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(["Task", "Category", "Worked at", "Progress", "Hours", "Status", "Accomplishment", "Blockers", "Next steps", "Remarks"])
    for entry in report_data["entries"]:
        for update in entry["updates"]:
            writer.writerow([
                entry["task"].get("subject", ""), entry["task"].get("category", ""),
                update.get("worked_at", "").isoformat() if update.get("worked_at") else "",
                update.get("progress", 0), update.get("hours_worked", 0), update.get("status", ""),
                update.get("accomplishment", ""), update.get("blockers", ""),
                update.get("next_steps", ""), update.get("remarks", ""),
            ])
    return Response(output.getvalue(), mimetype="text/csv", headers={"Content-Disposition": f"attachment; filename=taskr-{period}-journal.csv"})


def report_from_request():
    period = request.args.get("period", "weekly")
    custom_start = custom_end = None
    if period == "custom":
        try:
            custom_start = datetime.strptime(request.args["start"], "%Y-%m-%d")
            custom_end = datetime.strptime(request.args["end"], "%Y-%m-%d").replace(hour=23, minute=59, second=59)
        except (KeyError, ValueError):
            period = "weekly"
    return build_report(period, custom_start, custom_end, request.args.get("category"), request.args.get("status"), request.args.get("project"))


@journal.route("/export.pdf")
@login_required
def export_pdf():
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    output = StringIO()
    report_data = report_from_request()
    pdf_buffer = __import__("io").BytesIO()
    document = canvas.Canvas(pdf_buffer, pagesize=letter)
    document.setFont("Helvetica-Bold", 16)
    document.drawString(48, 750, "Taskr Work Journal")
    document.setFont("Helvetica", 10)
    document.drawString(48, 732, report_data["period_label"])
    y = 700
    for label, value in (("Tasks", report_data["stats"]["tasks"]), ("Updates", report_data["stats"]["updates"]), ("Hours", f'{report_data["stats"]["actual_hours"]:.2f}'), ("Accomplishments", report_data["stats"]["accomplishments"])):
        document.drawString(48, y, f"{label}: {value}")
        y -= 18
    y -= 10
    for entry in report_data["entries"]:
        for update in entry["updates"]:
            text = f'{entry["task"].get("subject", "")} | {update.get("worked_at", "")} | {update.get("progress", 0)}% | {update.get("hours_worked", 0)}h'
            document.drawString(48, y, text[:105])
            y -= 15
            if y < 50:
                document.showPage()
                y = 750
    document.save()
    return Response(pdf_buffer.getvalue(), mimetype="application/pdf", headers={"Content-Disposition": "attachment; filename=taskr-journal.pdf"})


@journal.route("/export.docx")
@login_required
def export_docx():
    from docx import Document
    report_data = report_from_request()
    document = Document()
    document.add_heading("Taskr Work Journal", 0)
    document.add_paragraph(report_data["period_label"])
    document.add_paragraph(f'Hours: {report_data["stats"]["actual_hours"]:.2f} | Updates: {report_data["stats"]["updates"]} | Accomplishments: {report_data["stats"]["accomplishments"]}')
    for entry in report_data["entries"]:
        document.add_heading(entry["task"].get("subject", "Task"), level=2)
        for update in entry["updates"]:
            document.add_paragraph(f'{update.get("worked_at", "")} | {update.get("progress", 0)}% | {update.get("hours_worked", 0)} hours')
            if update.get("accomplishment"):
                document.add_paragraph(update["accomplishment"])
    buffer = __import__("io").BytesIO()
    document.save(buffer)
    return Response(buffer.getvalue(), mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=taskr-journal.docx"})


@journal.route("/gantt")
@login_required
def gantt():
    history = []
    tasks = Task.all(owner_id=current_user.id)
    dates = [date for task in tasks for date in (task.get("start_date"), task.get("due_date")) if date]
    chart_start = min(dates) if dates else datetime.utcnow()
    chart_end = max(dates) if dates else chart_start + timedelta(days=1)
    span = max((chart_end - chart_start).total_seconds(), 1)
    axis = []
    for fraction in (0, .25, .5, .75, 1):
        axis.append({"position": fraction * 100, "label": (chart_start + timedelta(seconds=span * fraction)).strftime("%b %d, %Y")})
    for task in tasks:
        updates = TaskUpdate.for_task(task["_id"])
        history.append({"task": task, "updates": updates})
    return render_template("journal/gantt.html", history=history, axis=axis)
